from pathlib import Path

from docx import Document
from openpyxl import Workbook


ROOT = Path(__file__).resolve().parents[1]
TEMPLATES_DIR = ROOT / "templates"


HEADERS = [
    "#",
    "Application Code",
    "Title",
    "Author",
    "Language of translation",
    "Publisher",
    "Email To",
    "Email Cc",
    "Amount",
    "First Installment",
]

ROWS = [
    [1, "GL-001", "Moonlit Library", "Eleni D.", "Italian", "Aegean Books", "rights@aegean.example", "contracts@aegean.example", "5000", "2500"],
    [2, "GL-002", "Sea of Marble", "Nikos P.", "German", "Helios Press", "info@helios.example", "legal@helios.example", "4200", "2100"],
    [3, "GL-003", "Threads of Athena", "Maria K.", "French", "Attica Editions", "editor@attica.example", "accounts@attica.example", "6100", "3050"],
]


def create_workbook(path: Path) -> None:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "8TH PERIOD"

    for column_index, header in enumerate(HEADERS, start=1):
        worksheet.cell(row=2, column=column_index, value=header)

    for row_offset, values in enumerate(ROWS, start=3):
        for column_index, value in enumerate(values, start=1):
            worksheet.cell(row=row_offset, column=column_index, value=value)

    workbook.save(path)


def create_word_template(path: Path) -> None:
    document = Document()
    document.add_heading("GreekLit Contract Template", level=1)
    document.add_paragraph(
        "Application {{APPLICATION_CODE}} covers the work {{TITLE}} by {{AUTHOR}}."
    )
    document.add_paragraph(
        "This agreement concerns the {{LANGUAGE}} translation to be delivered by {{PUBLISHER}}."
    )
    document.add_paragraph(
        "The approved amount is {{AMOUNT}} EUR and the first installment is {{FIRST_INSTALLMENT}} EUR."
    )
    document.add_paragraph(
        "Please return the signed agreement and keep reference {{ID}} in your correspondence."
    )
    document.save(path)


def create_email_template(path: Path) -> None:
    path.write_text(
        "\n".join(
            [
                "Subject: GreekLit | {{APPLICATION_CODE}} | {{TITLE}}",
                "To: {{EMAIL_TO}}",
                "Cc: {{EMAIL_CC}}",
                "",
                "Dear {{PUBLISHER}},",
                "",
                "Please find attached the contract for {{TITLE}} by {{AUTHOR}}.",
                "The approved amount is {{AMOUNT}} EUR and the first installment is {{FIRST_INSTALLMENT}} EUR.",
                "",
                "Best regards,",
                "The GreekLit Team",
                "",
            ]
        ),
        encoding="utf-8",
    )


def main() -> None:
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    create_workbook(TEMPLATES_DIR / "starter-workbook.xlsx")
    create_word_template(TEMPLATES_DIR / "starter-contract-template.docx")
    create_email_template(TEMPLATES_DIR / "starter-email-template.txt")


if __name__ == "__main__":
    main()
